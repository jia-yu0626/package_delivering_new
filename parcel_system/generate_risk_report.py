# -*- coding: utf-8 -*-
"""
物流追蹤系統 - 風險評估報告產生器
生成 Word 格式的風險評估文件
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """設定儲存格背景顏色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_border(table):
    """設定表格邊框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

def create_risk_assessment_report():
    """建立風險評估報告 Word 文件"""
    doc = Document()
    
    # 設定文件樣式
    style = doc.styles['Normal']
    style.font.name = '微軟正黑體'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    style.font.size = Pt(12)
    
    # ====== 標題 ======
    title = doc.add_heading('物流追蹤系統風險評估報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副標題
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Parcel Tracking System Risk Assessment Report')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # ====== 專案資訊 ======
    doc.add_heading('一、專案資訊', level=1)
    
    info_table = doc.add_table(rows=4, cols=2)
    set_table_border(info_table)
    info_data = [
        ('專案名稱', '物流追蹤與計費系統 (Parcel Tracking System)'),
        ('技術架構', 'Python Flask + SQLAlchemy + SQLite'),
        ('評估日期', '2025年12月28日'),
        ('評估人員', '系統開發團隊'),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], 'E8E8E8')
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # ====== 風險評估矩陣說明 ======
    doc.add_heading('二、風險評估矩陣說明', level=1)
    
    matrix_para = doc.add_paragraph()
    matrix_para.add_run('風險等級判定標準：').bold = True
    
    legend_table = doc.add_table(rows=4, cols=3)
    set_table_border(legend_table)
    legend_data = [
        ('風險等級', '機率 × 影響', '說明'),
        ('🔴 高風險', '高×嚴重/災難 或 中度×災難', '需立即處理，優先分配資源'),
        ('🟡 中風險', '中度×嚴重 或 低×災難', '需制定應對計畫，持續監控'),
        ('🟢 低風險', '低×中度 或 中度×中度以下', '持續監控，必要時採取行動'),
    ]
    for i, (level, criteria, desc) in enumerate(legend_data):
        row = legend_table.rows[i]
        row.cells[0].text = level
        row.cells[1].text = criteria
        row.cells[2].text = desc
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '4472C4')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # ====== 專案通用風險 ======
    doc.add_heading('三、專案通用風險評估', level=1)
    
    # 風險表格
    risk_table = doc.add_table(rows=12, cols=5)
    set_table_border(risk_table)
    
    risks = [
        ('編號', '風險描述', '機率', '影響', '風險等級'),
        ('1', '低估開發軟體的所需時間 - 系統功能複雜（6種角色、多種業務流程），可能導致開發時程延誤', '高', '嚴重', '🔴 高風險'),
        ('4', '因為組織重整，改由不同的管理階層負責此專案 - 專案中途更換負責人導致方向改變', '高', '嚴重', '🔴 高風險'),
        ('5', '因組織的財務問題迫使專案預算遭刪減 - 資金不足影響開發資源', '低', '災難', '🟡 中風險'),
        ('6', '無法僱用到具有所需技能的人員 - Flask/SQLAlchemy 等技術棧需要專業人才', '高', '災難', '🔴 高風險'),
        ('7', '重要成員生病了，而且在關鍵時刻無法工作 - 關鍵開發人員無法工作影響進度', '中度', '嚴重', '🟡 中風險'),
        ('8', '無法對人員提供必要的訓練 - 用戶（倉儲人員、司機）不熟悉系統操作', '中度', '嚴重', '🟡 中風險'),
        ('9', '需求的變更導致主要設計需要重做 - 客戶要求變更計費規則、物流流程等核心邏輯', '中度', '嚴重', '🟡 中風險'),
        ('10', '客戶不瞭解需求變更的影響 - 利害關係人不理解改動成本', '中度', '中度', '🟢 低風險'),
        ('11', '系統所用資料庫的每秒交易量未能如預期的多 - SQLite 在高併發場景下可能成為瓶頸', '中度', '嚴重', '🟡 中風險'),
        ('12', '再利用的軟體元件有缺陷，必須先修復後才能再利用 - Flask 擴充套件或第三方函式庫有漏洞', '中度', '嚴重', '🟡 中風險'),
        ('14', '軟體工具無法整合在一起工作 - 前後端整合問題、API 相容性問題', '低', '中度', '🟢 低風險'),
    ]
    
    for i, row_data in enumerate(risks):
        row = risk_table.rows[i]
        for j, value in enumerate(row_data):
            row.cells[j].text = value
            if i == 0:
                set_cell_shading(row.cells[j], '4472C4')
                row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                row.cells[j].paragraphs[0].runs[0].bold = True
            elif '高風險' in str(row_data[4]):
                if j == 4:
                    set_cell_shading(row.cells[j], 'FFCCCC')
            elif '中風險' in str(row_data[4]):
                if j == 4:
                    set_cell_shading(row.cells[j], 'FFFFCC')
            elif '低風險' in str(row_data[4]):
                if j == 4:
                    set_cell_shading(row.cells[j], 'CCFFCC')
    
    doc.add_paragraph()
    
    # ====== 系統特定風險 ======
    doc.add_heading('四、系統特定風險評估', level=1)
    
    sys_risk_table = doc.add_table(rows=7, cols=5)
    set_table_border(sys_risk_table)
    
    sys_risks = [
        ('編號', '系統特定風險', '機率', '影響', '風險等級'),
        ('S1', '包裹追蹤資料遺失 - 追蹤事件記錄(TrackingEvent)未正確儲存', '低', '災難', '🟡 中風險'),
        ('S2', '計費錯誤 - PricingRule 計算邏輯錯誤導致金額不正確', '中度', '嚴重', '🟡 中風險'),
        ('S3', '權限控制繞過 - 不同角色(Customer/Driver/Admin)權限邊界模糊', '中度', '災難', '🔴 高風險'),
        ('S4', '密碼安全性不足 - 用戶帳號被盜取或暴力破解', '中度', '嚴重', '🟡 中風險'),
        ('S5', '系統無法處理尖峰流量 - 促銷活動期間包裹量暴增', '中度', '嚴重', '🟡 中風險'),
        ('S6', '司機指派邏輯失效 - auto_assign_packages() 無法正確分配包裹', '低', '中度', '🟢 低風險'),
    ]
    
    for i, row_data in enumerate(sys_risks):
        row = sys_risk_table.rows[i]
        for j, value in enumerate(row_data):
            row.cells[j].text = value
            if i == 0:
                set_cell_shading(row.cells[j], '4472C4')
                row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                row.cells[j].paragraphs[0].runs[0].bold = True
            elif '高風險' in str(row_data[4]):
                if j == 4:
                    set_cell_shading(row.cells[j], 'FFCCCC')
            elif '中風險' in str(row_data[4]):
                if j == 4:
                    set_cell_shading(row.cells[j], 'FFFFCC')
            elif '低風險' in str(row_data[4]):
                if j == 4:
                    set_cell_shading(row.cells[j], 'CCFFCC')
    
    doc.add_paragraph()
    
    # ====== 風險應對策略 ======
    doc.add_heading('五、風險應對策略', level=1)
    
    # 高風險
    doc.add_heading('🔴 高優先處理（高風險）', level=2)
    high_risks = [
        ('1', '低估開發時間', '建立實際的時程估算、採用敏捷迭代、分階段交付'),
        ('6', '缺乏技術人員', '確保團隊具備 Python/Flask 經驗、準備培訓計畫'),
        ('4', '管理層變動', '完善專案文件與交接程序、建立知識庫'),
        ('S3', '權限控制繞過', '加強安全性測試、程式碼審查、嚴格的角色驗證'),
    ]
    
    high_table = doc.add_table(rows=len(high_risks)+1, cols=3)
    set_table_border(high_table)
    header = high_table.rows[0]
    for j, h in enumerate(['編號', '風險', '應對策略']):
        header.cells[j].text = h
        set_cell_shading(header.cells[j], 'C00000')
        header.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        header.cells[j].paragraphs[0].runs[0].bold = True
    
    for i, (num, risk, strategy) in enumerate(high_risks):
        row = high_table.rows[i+1]
        row.cells[0].text = num
        row.cells[1].text = risk
        row.cells[2].text = strategy
    
    doc.add_paragraph()
    
    # 中風險
    doc.add_heading('🟡 中優先處理（中風險）', level=2)
    med_risks = [
        ('11', '資料庫效能', '規劃未來遷移至 PostgreSQL、建立效能監控'),
        ('S2', '計費錯誤', '增加計費邏輯的單元測試覆蓋率、帳單審核流程'),
        ('12', '第三方元件缺陷', '定期執行 pip audit 檢查漏洞、選用成熟套件'),
        ('7', '成員生病', '知識共享、程式碼審查制度、交叉培訓'),
        ('S4', '密碼安全性', '密碼雜湊(已實作)、登入失敗限制、考慮雙因素認證'),
    ]
    
    med_table = doc.add_table(rows=len(med_risks)+1, cols=3)
    set_table_border(med_table)
    header = med_table.rows[0]
    for j, h in enumerate(['編號', '風險', '應對策略']):
        header.cells[j].text = h
        set_cell_shading(header.cells[j], 'FFC000')
        header.cells[j].paragraphs[0].runs[0].bold = True
    
    for i, (num, risk, strategy) in enumerate(med_risks):
        row = med_table.rows[i+1]
        row.cells[0].text = num
        row.cells[1].text = risk
        row.cells[2].text = strategy
    
    doc.add_paragraph()
    
    # 低風險
    doc.add_heading('🟢 持續監控（低風險）', level=2)
    low_risks = [
        ('14', '整合問題', '維持現有 CI/CD 流程、自動化測試'),
        ('S6', '指派邏輯', '現有實作已可滿足需求、手動指派備案'),
        ('10', '客戶理解度', '建立變更影響評估流程、透明溝通'),
    ]
    
    low_table = doc.add_table(rows=len(low_risks)+1, cols=3)
    set_table_border(low_table)
    header = low_table.rows[0]
    for j, h in enumerate(['編號', '風險', '應對策略']):
        header.cells[j].text = h
        set_cell_shading(header.cells[j], '70AD47')
        header.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        header.cells[j].paragraphs[0].runs[0].bold = True
    
    for i, (num, risk, strategy) in enumerate(low_risks):
        row = low_table.rows[i+1]
        row.cells[0].text = num
        row.cells[1].text = risk
        row.cells[2].text = strategy
    
    doc.add_paragraph()
    
    # ====== 結論 ======
    doc.add_heading('六、結論與建議', level=1)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run('風險總結：\n').bold = True
    conclusion.add_run('本系統共識別 17 項風險，其中高風險 4 項、中風險 10 項、低風險 3 項。\n\n')
    
    conclusion.add_run('主要建議：\n').bold = True
    conclusion.add_run('1. 優先處理權限控制與安全性相關風險\n')
    conclusion.add_run('2. 建立完善的測試與品質保證流程\n')
    conclusion.add_run('3. 制定詳細的專案時程與人力規劃\n')
    conclusion.add_run('4. 建立知識管理與文件化制度\n')
    conclusion.add_run('5. 定期進行風險審查與更新\n')
    
    # 儲存文件
    output_path = os.path.join(os.path.dirname(__file__), '物流追蹤系統_風險評估報告.docx')
    doc.save(output_path)
    print(f"[OK] Risk Assessment Report Generated: {output_path}")
    return output_path

if __name__ == '__main__':
    create_risk_assessment_report()
